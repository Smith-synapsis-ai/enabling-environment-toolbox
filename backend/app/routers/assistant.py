"""Assistant router — conversational challenge assistant (Task A6).

Exposes:

* ``WS /ws/challenge`` — streams orchestrator events for a challenge turn.
  The client sends JSON text frames of the shape::

      {"challenge_text": "...", "session_id": "<uuid>"}

  For each message the server iterates ``run_challenge`` and forwards every
  event dict as a single JSON text frame, then sends a
  ``{"type": "turn_complete"}`` sentinel.  The socket stays open so the same
  connection can carry multiple turns (refinement / approval).

* ``GET /api/assistant/sessions`` — lists persisted report-draft sessions.

* ``GET /api/assistant/sessions/{session_id}/draft`` — returns the parsed
  draft JSON plus a ``rendered_markdown`` field.

Concurrency note
----------------
The agent layer binds the active report session via a module-level
``set_current_report_session`` call inside ``run_challenge``, i.e. the
orchestrator can only safely run ONE session at a time per process.  We
therefore serialize all turns through a single global ``asyncio.Lock``
(across all websocket connections).  This is a deliberate, documented
simplification for the prototype: concurrent users queue rather than
corrupt each other's report state.
"""

from __future__ import annotations

import asyncio
import json
import logging
import uuid

import re

from fastapi import APIRouter, HTTPException, Query, WebSocket, WebSocketDisconnect, UploadFile, File
from fastapi.responses import StreamingResponse
import io

# Importing agents.orchestrator also swaps the default in-memory report store
# for the SQLite-backed one (side effect at module import).  run_challenge is
# referenced lazily (module attribute) in the WS handler so tests can
# monkeypatch ``app.routers.assistant.run_challenge``.
from agents.orchestrator import run_challenge  # noqa: F401
from agents.report_state import ReportDraft, get_report_store
from persistence.store import SqliteReportStore

logger = logging.getLogger("ee.assistant")

router = APIRouter()

# Global lock: the orchestrator supports one active session per process.
_CHALLENGE_LOCK = asyncio.Lock()


def _is_uuid(value: str) -> bool:
    try:
        uuid.UUID(value)
        return True
    except (ValueError, AttributeError, TypeError):
        return False


@router.websocket("/ws/challenge")
async def challenge_ws(websocket: WebSocket) -> None:
    """Multi-turn challenge socket.

    Each inbound JSON message triggers one orchestrator turn.  Events are
    forwarded verbatim; a ``turn_complete`` sentinel marks the end of each
    turn (sent even after errors so the client can unblock its input).
    """
    await websocket.accept()
    try:
        while True:
            raw = await websocket.receive_text()

            # -- parse + validate the inbound message -------------------
            try:
                payload = json.loads(raw)
            except json.JSONDecodeError:
                await websocket.send_text(json.dumps(
                    {"type": "error", "message": "Invalid JSON message."}
                ))
                await websocket.send_text(json.dumps({"type": "turn_complete"}))
                continue

            challenge_text = str(payload.get("challenge_text") or "").strip()
            session_id = payload.get("session_id")

            if not challenge_text:
                await websocket.send_text(json.dumps(
                    {"type": "error", "message": "challenge_text is required."}
                ))
                await websocket.send_text(json.dumps({"type": "turn_complete"}))
                continue

            # The orchestrator silently regenerates non-UUID session ids,
            # which would break resume — fail fast instead.
            if session_id is not None and not _is_uuid(str(session_id)):
                await websocket.send_text(json.dumps({
                    "type": "error",
                    "message": (
                        "session_id must be a valid UUID "
                        f"(got {session_id!r})."
                    ),
                }))
                await websocket.send_text(json.dumps({"type": "turn_complete"}))
                continue

            # -- run one orchestrator turn (serialized globally) --------
            try:
                async with _CHALLENGE_LOCK:
                    # 3600s hard timeout INSIDE the lock so lock-wait time is
                    # not charged against the user's turn.  This matches the
                    # ALB idle_timeout (also 3600s) so ALB never drops the
                    # connection before the agent can finish a long turn.
                    # The earlier 220s figure was aspirational only; the real
                    # constraint is total pipeline time for complex challenges.
                    async with asyncio.timeout(3600):
                        async for event in run_challenge(
                            challenge_text,
                            session_id=str(session_id) if session_id else None,
                        ):
                            await websocket.send_text(
                                json.dumps(event, default=str)
                            )
            except WebSocketDisconnect:
                raise
            except TimeoutError:
                # asyncio.TimeoutError (== TimeoutError in 3.11+) — lock is
                # released automatically when the context manager exits.
                logger.warning(
                    "Challenge turn timed out after 220s (session=%s)", session_id
                )
                try:
                    await websocket.send_text(json.dumps({
                        "type": "error",
                        "code": "timeout",
                        "message": "Challenge timed out — please try again.",
                    }))
                except Exception:
                    raise WebSocketDisconnect()
            except Exception as exc:  # noqa: BLE001 — never crash the socket
                logger.exception("Challenge turn failed")
                try:
                    await websocket.send_text(json.dumps({
                        "type": "error",
                        "message": f"Turn failed: {exc}",
                    }))
                except Exception:  # socket already gone
                    raise WebSocketDisconnect() from exc

            await websocket.send_text(json.dumps({"type": "turn_complete"}))
    except WebSocketDisconnect:
        logger.info("Challenge websocket disconnected")


@router.get("/api/assistant/sessions")
async def list_sessions() -> dict:
    """List persisted report-draft sessions (most recently updated first)."""
    store = get_report_store()
    sessions = await store.list_sessions()
    return {"sessions": sessions}


@router.get("/api/assistant/sessions/{session_id}/draft")
async def get_session_draft(session_id: str) -> dict:
    """Return the parsed report draft plus rendered markdown, or 404."""
    store = get_report_store()
    raw = await store.load_draft(session_id)
    if raw is None:
        raise HTTPException(status_code=404, detail="No draft for this session")
    draft = ReportDraft.from_json(raw)
    data = json.loads(draft.to_json())
    data["rendered_markdown"] = draft.render_markdown()
    return data


# ---------------------------------------------------------------------------
# Admin endpoints (session management)
# ---------------------------------------------------------------------------

@router.delete("/api/assistant/sessions/{session_id}")
async def delete_session(session_id: str) -> dict:
    """Delete the report draft for a session (admin / test cleanup).

    Returns 200 with ``{"deleted": true}`` if the session existed, or 404 if
    it did not.  The WebSocket connection (if any) is NOT forcibly closed; the
    client will receive a missing-draft error on its next draft fetch.
    """
    store = get_report_store()
    # delete_session / reap_stale_sessions are SqliteReportStore extensions
    # not declared on the base ReportStore Protocol (Protocol is frozen per
    # wave3-interfaces.md §1 contract).  The production store is always
    # SqliteReportStore (set at orchestrator import), so this guard is a
    # safety net and type-narrowing hint, not a real code path.
    if not isinstance(store, SqliteReportStore):
        raise HTTPException(
            status_code=501,
            detail="Session admin endpoints require the SQLite report store.",
        )
    deleted = await store.delete_session(session_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="No draft for this session")
    logger.info("Session %s deleted via admin endpoint", session_id)
    return {"deleted": True, "session_id": session_id}


@router.post("/api/admin/sessions/reap")
async def reap_sessions(max_age_hours: int = 24) -> dict:
    """Delete report drafts not updated in the last *max_age_hours* hours.

    Query parameter: ``max_age_hours`` (default: 24).  Returns a count of
    deleted sessions.  Intended for manual maintenance or a scheduled cron.
    Minimum value is 1 hour (values < 1 are clamped to 1).
    """
    max_age_hours = max(1, max_age_hours)
    store = get_report_store()
    if not isinstance(store, SqliteReportStore):
        raise HTTPException(
            status_code=501,
            detail="Session admin endpoints require the SQLite report store.",
        )
    deleted = await store.reap_stale_sessions(max_age_hours=max_age_hours)
    logger.info("Session reap: deleted %d drafts older than %dh", deleted, max_age_hours)
    return {"deleted_count": deleted, "max_age_hours": max_age_hours}


@router.post("/api/assistant/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    Accept a file upload, extract its text content, and return it for injection
    into the agent conversation context.
    Supports: .txt, .md, .csv, .json, .py, .js, .ts, .html, .xml (plain text)
    Returns: { "filename": str, "content": str, "char_count": int }
    """
    MAX_SIZE = 500_000  # 500 KB limit to avoid context bloat

    content_bytes = await file.read()
    if len(content_bytes) > MAX_SIZE:
        return {
            "filename": file.filename,
            "content": f"[File too large to process — {len(content_bytes):,} bytes, max {MAX_SIZE:,} bytes. Please paste the relevant excerpt directly into the chat.]",
            "char_count": 0,
        }

    # Attempt UTF-8 decode; fall back to latin-1 for binary-ish files
    try:
        text = content_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content_bytes.decode("latin-1")
        except Exception:
            return {
                "filename": file.filename,
                "content": f"[Binary file '{file.filename}' cannot be extracted as text. Please convert to plain text or paste the relevant content directly.]",
                "char_count": 0,
            }

    return {
        "filename": file.filename,
        "content": text,
        "char_count": len(text),
    }


# ---------------------------------------------------------------------------
# Report export endpoints (PDF + Word)
# ---------------------------------------------------------------------------

@router.get("/api/assistant/sessions/{session_id}/draft/export")
async def export_draft(
    session_id: str,
    format: str = Query("pdf", pattern="^(pdf|docx)$"),
):
    """
    Export the session's report draft as a PDF or Word document.
    Query param: ?format=pdf (default) or ?format=docx
    Returns a streaming file download.
    """
    store = get_report_store()
    raw = await store.load_draft(session_id)
    if raw is None:
        raise HTTPException(status_code=404, detail="No draft for this session")
    draft = ReportDraft.from_json(raw)
    markdown_text = draft.render_markdown()

    # Extract a title from the first H1 or H2 line, or use default
    title_match = re.search(r'^#{1,2}\s+(.+)$', markdown_text, re.MULTILINE)
    doc_title = title_match.group(1).strip() if title_match else "Scaling Challenge Report"

    if format == "docx":
        return await _export_docx(session_id, doc_title, markdown_text, draft)
    else:
        return await _export_pdf(session_id, doc_title, markdown_text, draft)


async def _export_docx(session_id: str, title: str, markdown_text: str, draft) -> StreamingResponse:
    """Generate a .docx file from the draft using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ── CGIAR header block ────────────────────────────────────────────────
    cgiar_green = RGBColor(3, 53, 41)   # #033529
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("CGIAR Enabling Environment Toolbox")
    header_run.font.color.rgb = cgiar_green
    header_run.font.size = Pt(10)
    header_run.font.bold = True
    header_para.add_run(f"\nScaling Challenge Report · Revision {draft.revision}")
    header_para.paragraph_format.space_after = Pt(6)

    # ── Title ─────────────────────────────────────────────────────────────
    title_para = doc.add_heading(title, level=1)
    title_para.runs[0].font.color.rgb = cgiar_green

    # ── Candidate tools table ─────────────────────────────────────────────
    candidate_tools = []
    try:
        import json as _json
        data = _json.loads(draft.to_json())
        candidate_tools = data.get("candidate_tools", [])
    except Exception:
        pass

    if candidate_tools:
        doc.add_heading("Recommended Tools", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Shading Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Tool"
        hdr_cells[1].text = "Status"
        for t in candidate_tools:
            row_cells = table.add_row().cells
            row_cells[0].text = t.get("title", "")
            row_cells[1].text = t.get("status", "candidate").title()
        doc.add_paragraph()  # spacing after table

    # ── Body: parse markdown into docx paragraphs ─────────────────────────
    for line in markdown_text.split('\n'):
        line = line.rstrip()
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = cgiar_green
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line == '' or line == '---':
            doc.add_paragraph()
        else:
            # Strip inline markdown formatting (bold, italic, links)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            if clean.strip():
                doc.add_paragraph(clean)

    # ── Footer ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph(f"Generated by EE Toolbox Assistant · Session {session_id[:8]}…")
    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"scaling-report-r{draft.revision}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def _export_pdf(session_id: str, title: str, markdown_text: str, draft) -> StreamingResponse:
    """Generate a PDF from the draft using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab not installed")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5*cm,
        rightMargin=2.5*cm,
        topMargin=2.5*cm,
        bottomMargin=2.5*cm,
    )

    # ── Style definitions ─────────────────────────────────────────────────
    cgiar_green = colors.HexColor('#033529')
    cgiar_accent = colors.HexColor('#4ade80')
    styles = getSampleStyleSheet()

    style_h1 = ParagraphStyle(
        'H1', parent=styles['Heading1'],
        textColor=cgiar_green, fontSize=18, spaceAfter=12,
    )
    style_h2 = ParagraphStyle(
        'H2', parent=styles['Heading2'],
        textColor=cgiar_green, fontSize=14, spaceAfter=8, spaceBefore=12,
    )
    style_h3 = ParagraphStyle(
        'H3', parent=styles['Heading3'],
        textColor=colors.HexColor('#1a5c40'), fontSize=12, spaceAfter=6, spaceBefore=8,
    )
    style_body = ParagraphStyle(
        'Body', parent=styles['Normal'],
        fontSize=10, spaceAfter=6, leading=16,
    )
    style_bullet = ParagraphStyle(
        'Bullet', parent=styles['Normal'],
        fontSize=10, spaceAfter=4, leftIndent=12, bulletIndent=4, leading=14,
    )
    style_caption = ParagraphStyle(
        'Caption', parent=styles['Normal'],
        fontSize=8, textColor=colors.grey, alignment=TA_CENTER,
    )
    style_header_brand = ParagraphStyle(
        'Brand', parent=styles['Normal'],
        textColor=colors.white, fontSize=9, leading=13,
    )

    story = []

    # ── Branded header banner ──────────────────────────────────────────────
    header_table = Table(
        [[Paragraph(
            f'<b>CGIAR Enabling Environment Toolbox</b><br/>'
            f'Scaling Challenge Report · Revision {draft.revision}',
            style_header_brand,
        )]],
        colWidths=['100%'],
    )
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), cgiar_green),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('LEFTPADDING', (0, 0), (-1, -1), 14),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 0.5*cm))

    # ── Document title ─────────────────────────────────────────────────────
    story.append(Paragraph(title, style_h1))
    story.append(HRFlowable(width='100%', thickness=1, color=cgiar_accent, spaceAfter=10))

    # ── Candidate tools table ──────────────────────────────────────────────
    candidate_tools = []
    try:
        import json as _json
        data = _json.loads(draft.to_json())
        candidate_tools = data.get("candidate_tools", [])
    except Exception:
        pass

    if candidate_tools:
        story.append(Paragraph("Recommended Tools", style_h2))
        tdata = [['Tool', 'Status']]
        for t in candidate_tools:
            tdata.append([t.get('title', ''), t.get('status', 'candidate').title()])
        tbl = Table(tdata, colWidths=[12*cm, 3*cm])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), cgiar_green),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0faf4')]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#d1fae5')),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.4*cm))

    # ── Body: parse markdown into reportlab flowables ──────────────────────
    for line in markdown_text.split('\n'):
        line = line.rstrip()
        # Strip inline markdown: links [text](url) → text
        clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean)
        clean = re.sub(r'\*(.+?)\*', r'<i>\1</i>', clean)

        if clean.startswith('### '):
            story.append(Paragraph(clean[4:], style_h3))
        elif clean.startswith('## '):
            story.append(Paragraph(clean[3:], style_h2))
        elif clean.startswith('# '):
            story.append(Paragraph(clean[2:], style_h1))
        elif line.startswith('- ') or line.startswith('* '):
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line[2:])
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            story.append(Paragraph(f'• {text}', style_bullet))
        elif line == '---':
            story.append(HRFlowable(width='100%', thickness=0.5, color=colors.grey, spaceAfter=6))
        elif not line.strip():
            story.append(Spacer(1, 0.2*cm))
        else:
            try:
                story.append(Paragraph(clean, style_body))
            except Exception:
                story.append(Paragraph(re.sub(r'<[^>]+>', '', clean), style_body))

    # ── Footer caption ─────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width='100%', thickness=0.5, color=colors.grey))
    story.append(Paragraph(
        f'Generated by EE Toolbox Assistant · Session {session_id[:8]}… · cgiar.org',
        style_caption,
    ))

    doc.build(story)
    buf.seek(0)

    filename = f"scaling-report-r{draft.revision}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
